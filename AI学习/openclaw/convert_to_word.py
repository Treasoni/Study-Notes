#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OpenClaw 数字人学习笔记 - 美化版
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# Wd_Table_Alignment not needed, removed

def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()

def set_cell_background(cell, color):
    """设置单元格背景色"""
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_notebook_doc():
    """创建学习笔记文档"""
    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.page_width = Cm(21)  # A4
    section.page_height = Cm(29.7)

    return doc

def add_cover(doc):
    """添加封面"""
    # 标题
    title = doc.add_paragraph('OpenClaw 数字人')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 64, 128)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_paragraph('学习笔记与实战指南')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for _ in range(5):
        doc.add_paragraph('')

    # 信息框
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = (
        f'📅 整理时间：2026年2月25日\n'
        f'📚 资料来源：阿里云开发者社区、CSDN、掘金、官方文档\n'
        f'🎯 核心问题：如何用 OpenClaw 做数字人？'
    )
    run = info.add_run(info_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    add_page_break(doc)

def add_section_title(doc, title, level=1):
    """添加章节标题"""
    p = doc.add_paragraph()
    run = p.add_run(title)

    colors = {
        1: (RGBColor(0, 64, 128), Pt(20)),
        2: (RGBColor(0, 96, 160), Pt(16)),
        3: (RGBColor(0, 128, 192), Pt(14)),
    }

    color, size = colors.get(level, colors[3])
    run.font.color.rgb = color
    run.font.size = size
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)

    return p

def add_note_box(doc, text, box_type='info'):
    """添加笔记框"""
    colors = {
        'info': {'bg': 'E6F3FF', 'border': '4A90E2', 'icon': '💡'},
        'warning': {'bg': 'FFF4E6', 'border': 'FF9500', 'icon': '⚠️'},
        'success': {'bg': 'E8F5E9', 'border': '4CAF50', 'icon': '✅'},
        'error': {'bg': 'FFEBEE', 'border': 'F44336', 'icon': '❌'},
        'tip': {'bg': 'F3E5F5', 'border': '9C27B0', 'icon': '💡'},
        'key': {'bg': 'FFF9C4', 'border': 'F57C00', 'icon': '🔑'},
    }

    style = colors.get(box_type, colors['info'])

    # 创建表格作为笔记框
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]

    # 设置背景色
    set_cell_background(cell, style['bg'])

    # 添加内容
    p = cell.paragraphs[0]
    run = p.add_run(f"{style['icon']} {text}")
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 设置边框（简化版，不设置复杂边框）

    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

    return table

def add_paragraph(doc, text, bold=False, indent=0):
    """添加普通段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.color.rgb = RGBColor(40, 40, 40)

    if indent > 0:
        p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)

    return p

def add_bullet(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(2)

    return p

def add_comparison_table(doc, headers, rows):
    """添加对比表格"""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_background(cell, '4472C4')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # 数据行
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    table.autofit = True
    return table

def main():
    """主函数"""
    doc = create_notebook_doc()

    # 封面
    add_cover(doc)

    # 目录页
    add_section_title(doc, '📑 笔记目录', 2)
    add_paragraph(doc, '01. OpenClaw 核心概念')
    add_paragraph(doc, '02. 部署方式选择')
    add_paragraph(doc, '03. 成本分析对比')
    add_paragraph(doc, '04. 实战应用场景')
    add_paragraph(doc, '05. 快速启动指南')
    add_paragraph(doc, '06. 常见问题解答')
    add_page_break(doc)

    # ===== 第一部分：核心概念 =====
    add_section_title(doc, '01 OpenClaw 核心概念', 1)

    add_note_box(doc, 'OpenClaw 不是虚拟形象的"数字人"，而是能执行实际任务的 AI 智能体', 'key')

    add_section_title(doc, '1.1 什么是 OpenClaw？', 2)

    add_paragraph(doc, 'OpenClaw 是一个 AI 网关和服务编排平台，可以理解为：')
    add_bullet(doc, '"大脑"：大语言模型（云端 API 或本地模型）')
    add_bullet(doc, '"身体"：OpenClaw 网关（部署在本地/云端）')
    add_bullet(doc, '"手脚"：各种技能（文件操作、浏览器控制、终端命令等）')

    add_section_title(doc, '1.2 核心能力', 2)
    add_bullet(doc, '任务执行：通过代码调用系统资源')
    add_bullet(doc, '决策推理：基于大模型的智能决策')
    add_bullet(doc, '多工具集成：连接各种 API 和服务')
    add_bullet(doc, '持续运行：7×24 小时待命')
    add_bullet(doc, '记忆学习：通过 MEMORY.md 构建长期记忆')

    add_note_box(doc, 'Slogan: "The AI that actually does things" —— 不仅仅聊天，而是真正做事', 'tip')

    add_page_break(doc)

    # ===== 第二部分：部署方式选择 =====
    add_section_title(doc, '02 部署方式选择', 1)

    add_section_title(doc, '2.1 三种部署方式', 2)

    add_note_box(doc, '新手推荐：官方一键脚本 + DeepSeek API（成本低、见效快）', 'info')

    # 部署方式对比表
    headers = ['方式', '难度', '成本', '适用场景']
    rows = [
        ['一键脚本', '⭐ 简单', 'API 费用 ¥5-50/月', '个人学习、快速测试'],
        ['npm/pnpm', '⭐⭐ 中等', 'API 费用 ¥5-50/月', '开发者、自定义需求'],
        ['Docker', '⭐⭐ 中等', '服务器 ¥100-300/月 + API', '生产环境、团队协作'],
        ['本地部署', '⭐⭐⭐ 复杂', '硬件 ¥5000-15000 + 电费', '隐私敏感、高频使用'],
    ]
    add_comparison_table(doc, headers, rows)

    add_section_title(doc, '2.2 AI 模型选择', 2)

    add_paragraph(doc, '云端 API（推荐新手）：')
    add_bullet(doc, 'DeepSeek：¥1-2/百万 tokens（性价比最高）')
    add_bullet(doc, '通义千问：¥2-6/百万 tokens（中文友好）')
    add_bullet(doc, 'Claude 3.5：¥25-75/百万 tokens（质量最高）')

    add_paragraph(doc, '本地模型（需要硬件）：')
    add_bullet(doc, 'Llama 3 8B：需要 16GB 显存')
    add_bullet(doc, 'Qwen 2 72B：需要 40GB+ 显存')
    add_bullet(doc, '通过 Ollama 本地运行')

    add_page_break(doc)

    # ===== 第三部分：成本分析 =====
    add_section_title(doc, '03 成本分析对比', 1)

    add_section_title(doc, '3.1 月度成本对比', 2)

    headers = ['使用量级', 'DeepSeek', '通义千问', '本地部署']
    rows = [
        ['轻度（50万 tokens）', '¥2', '¥5', '¥0（需硬件投入）'],
        ['中度（200万 tokens）', '¥5', '¥20', '¥0（需硬件投入）'],
        ['重度（500万 tokens）', '¥15', '¥60', '¥0（需硬件投入）'],
        ['企业级（2000万）', '¥50', '¥240', '¥0（需硬件投入）'],
    ]
    add_comparison_table(doc, headers, rows)

    add_section_title(doc, '3.2 本地部署投资回报分析', 2)

    add_note_box(doc, '关键结论：个人/小团队用云端 API 更划算；大团队或隐私场景本地部署更合适', 'key')

    add_paragraph(doc, '本地部署成本：')
    add_bullet(doc, '硬件投入：¥8,000-15,000（16GB 显存 GPU）')
    add_bullet(doc, '电费：¥100-200/月（24小时运行）')
    add_bullet(doc, '回本周期：约 38 个月（相比 DeepSeek）')

    add_paragraph(doc, '推荐方案：')
    add_bullet(doc, '测试/学习：DeepSeek API（月费 ¥5-15）')
    add_bullet(doc, '小团队：DeepSeek + 备用 Claude（月费 ¥50-100）')
    add_bullet(doc, '大企业：混合部署（本地 70% + 云端 30%）')

    add_page_break(doc)

    # ===== 第四部分：实战应用场景 =====
    add_section_title(doc, '04 实战应用场景', 1)

    add_section_title(doc, '4.1 个人生产力', 2)

    add_paragraph(doc, '办公自动化（效率提升 5-30 倍）：')
    add_bullet(doc, '邮件自动整理：30分钟 → 5分钟')
    add_bullet(doc, '会议纪要生成：会议结束即有纪要')
    add_bullet(doc, '发票信息录入：拍照自动识别，准确率 95%+')
    add_bullet(doc, '文档格式转换：100份文档2分钟完成')
    add_bullet(doc, '报告自动生成：周报月报一键生成')

    add_section_title(doc, '4.2 团队协作', 2)

    add_paragraph(doc, '企业场景：')
    add_bullet(doc, '飞书/钉钉智能客服：80% 问题自动回答')
    add_bullet(doc, '工单自动分类：分配准确率 90%')
    add_bullet(doc, '项目进度汇报：每日自动生成')
    add_bullet(doc, '跨时区会议协调：自动选择最佳时间')

    add_section_title(doc, '4.3 开发者工具', 2)

    add_paragraph(doc, '技术场景：')
    add_bullet(doc, '代码审查：发现 80% 潜在 bug')
    add_bullet(doc, '单元测试生成：覆盖率提升到 90%')
    add_bullet(doc, '自动化部署：一键上线')
    add_bullet(doc, '日志分析：异常自动发现')

    add_note_box(doc, '总计 98 个真实案例，涵盖办公、协作、开发、金融、教育等多个领域', 'success')

    add_page_break(doc)

    # ===== 第五部分：快速启动指南 =====
    add_section_title(doc, '05 快速启动指南', 1)

    add_section_title(doc, '5.1 Day 0 - 前期准备', 2)

    add_bullet(doc, '✅ 明确使用场景和目标')
    add_bullet(doc, '✅ 确定用户规模和预算')
    add_bullet(doc, '✅ 安装 Node.js ≥ 22')
    add_bullet(doc, '✅ 准备服务器或本地机器')

    add_section_title(doc, '5.2 Day 1 - 基础部署', 2)

    add_paragraph(doc, '1. 安装 OpenClaw：')
    add_paragraph(doc, 'curl -fsSL https://openclaw.bot/install.sh | bash', bold=False, indent=0.3)

    add_paragraph(doc, '2. 获取 API Key（推荐 DeepSeek）：')
    add_bullet(doc, '访问 https://www.deepseek.com 注册')
    add_bullet(doc, '创建 API Key')

    add_paragraph(doc, '3. 配置并启动：')
    add_paragraph(doc, 'openclaw gateway', bold=False, indent=0.3)

    add_section_title(doc, '5.3 Day 2-3 - 功能配置', 2)

    add_bullet(doc, '连接通讯渠道（WhatsApp/Telegram/飞书等）')
    add_bullet(doc, '启用基础技能（文件管理、浏览器自动化）')
    add_bullet(doc, '构建 MEMORY.md 知识库')
    add_bullet(doc, '测试基本对话功能')

    add_note_box(doc, '详细步骤见《OpenClaw安装教程.md》', 'info')

    add_page_break(doc)

    # ===== 第六部分：常见问题 =====
    add_section_title(doc, '06 常见问题解答', 1)

    add_section_title(doc, '6.1 核心问题', 2)

    add_paragraph(doc, 'Q1: OpenClaw 和 ChatGPT 有什么区别？')
    add_paragraph(doc, 'A: ChatGPT 只能聊天，OpenClaw 能执行实际操作（发邮件、操作文件等）。可以把 OpenClaw 理解为"给 ChatGPT 加上了手脚"。', indent=0.2)

    add_paragraph(doc, 'Q2: 没有编程基础能用吗？')
    add_paragraph(doc, 'A: 可以。OpenClaw 提供一键安装、可视化界面和预置技能。但自定义技能需要编程知识。', indent=0.2)

    add_paragraph(doc, 'Q3: 本地部署需要什么配置？')
    add_paragraph(doc, 'A: 最低 16GB 内存 + 8GB 显存 GPU。推荐 16GB 显存（如 RTX 4080）。', indent=0.2)

    add_section_title(doc, '6.2 安全问题', 2)

    add_note_box(doc, '⚠️ OpenClaw 具有系统权限，需要谨慎配置安全策略', 'warning')

    add_bullet(doc, 'API Key 安全：加密存储、定期轮换')
    add_bullet(doc, '权限控制：命令白名单、文件访问限制')
    add_bullet(doc, '审计日志：记录所有操作')
    add_bullet(doc, '敏感数据：本地处理或脱敏')

    add_page_break(doc)

    # ===== 第七部分：学习资源 =====
    add_section_title(doc, '07 学习资源', 1)

    add_section_title(doc, '7.1 官方资源', 2)

    add_bullet(doc, '官方网站：https://openclaw.ai')
    add_bullet(doc, '官方文档：https://docs.openclaw.ai')
    add_bullet(doc, 'GitHub：https://github.com/openclaw/openclaw')
    add_bullet(doc, '中文社区：https://www.moltcn.com')

    add_section_title(doc, '7.2 推荐学习路径', 2)

    add_paragraph(doc, '第 1-2 周：基础入门')
    add_bullet(doc, '阅读官方文档', level=1)
    add_bullet(doc, '完成快速开始', level=1)
    add_bullet(doc, '部署第一个机器人', level=1)

    add_paragraph(doc, '第 2-4 周：进阶应用')
    add_bullet(doc, '学习技能开发', level=1)
    add_bullet(doc, '配置多种通讯渠道', level=1)
    add_bullet(doc, '构建知识库', level=1)

    add_paragraph(doc, '第 1-3 月：高级定制')
    add_bullet(doc, '自定义技能开发', level=1)
    add_bullet(doc, '性能优化', level=1)
    add_bullet(doc, '安全加固', level=1)

    add_page_break(doc)

    # ===== 第八部分：决策树 =====
    add_section_title(doc, '08 决策参考', 1)

    add_section_title(doc, '8.1 快速决策表', 2)

    add_note_box(doc, '选择部署方式的三要素：隐私要求、使用频率、团队规模', 'key')

    headers = ['场景', '推荐方案', '月成本', '技术要求']
    rows = [
        ['个人学习', 'DeepSeek API', '¥5-15', '低'],
        ['个人重度', '本地 GPU', '¥100-200 电费', '中'],
        ['小团队（2-5人）', '云端 API', '¥50-100', '中'],
        ['中团队（5-20人）', '云端 API', '¥100-300', '中'],
        ['大企业/隐私', '本地部署', '硬件投入', '高'],
    ]
    add_comparison_table(doc, headers, rows)

    add_section_title(doc, '8.2 关键决策点', 2)

    add_paragraph(doc, '选择本地部署 if：')
    add_bullet(doc, '数据隐私要求极高（金融、医疗）')
    add_bullet(doc, '需要离线运行')
    add_bullet(doc, '长期高频使用（月成本 > ¥500）')
    add_bullet(doc, '有技术维护团队')

    add_paragraph(doc, '选择云端 API if：')
    add_bullet(doc, '快速验证想法')
    add_bullet(doc, '预算有限（< ¥200/月）')
    add_bullet(doc, '弹性业务需求')
    add_bullet(doc, '无维护团队')

    # ===== 结尾 =====
    add_page_break(doc)
    add_section_title(doc, '📝 总结', 1)

    add_note_box(doc, 'OpenClaw 是一个强大的 AI 智能体平台，核心价值在于"执行任务"而非"聊天对话"', 'success')

    add_paragraph(doc, '核心要点：')
    add_bullet(doc, '概念：OpenClaw = AI 大脑 + 执行手脚')
    add_bullet(doc, '部署：新手用一键脚本 + DeepSeek API')
    add_bullet(doc, '成本：个人月费 ¥5-50，本地部署需 ¥8000+ 硬件')
    add_bullet(doc, '场景：98 个真实案例，覆盖办公、协作、开发等领域')
    add_bullet(doc, '安全：注意权限控制和 API Key 保护')

    add_paragraph(doc, '')
    add_paragraph(doc, '下一步行动：')
    add_bullet(doc, '1. 明确你想解决的具体问题')
    add_bullet(doc, '2. 选择合适的部署方式（参考决策表）')
    add_bullet(doc, '3. 按照快速启动指南动手实践')
    add_bullet(doc, '4. 加入社区交流学习')

    add_paragraph(doc, '')
    add_paragraph(doc, '祝您成功构建自己的数字员工！🚀', bold=True)

    # 保存
    output_file = '/Users/zhqznc/Documents/项目/AI学习/openclaw/OpenClaw学习笔记.docx'
    doc.save(output_file)

    print('✅ 学习笔记已生成')
    print(f'📄 文件位置：{output_file}')
    print(f'📊 文件大小：{(Path(output_file).stat().st_size / 1024):.1f} KB')
    print('')
    print('📚 笔记内容：')
    print('   ✅ 封面页')
    print('   ✅ 目录页')
    print('   ✅ 8 个核心章节')
    print('   ✅ 彩色提示框')
    print('   ✅ 对比表格')
    print('   ✅ 快速启动指南')
    print('   ✅ 决策参考表')
    print('   ✅ FAQ 解答')
    print('')
    print('💡 提示：在 Word 中打开即可查看完整格式')

if __name__ == '__main__':
    from pathlib import Path
    main()
